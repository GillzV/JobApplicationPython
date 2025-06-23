import re
import PyPDF2
from docx import Document
import json
from typing import Dict, List, Any, Optional, Tuple
import logging
from datetime import datetime
import os

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ResumeParserV2:
    def __init__(self):
        """Initialize the improved resume parser."""
        
        # Define section headers with multiple variations
        self.section_patterns = {
            'contact': [
                r'contact\s*information',
                r'personal\s*details',
                r'personal\s*information',
                r'contact\s*details'
            ],
            'summary': [
                r'professional\s*summary',
                r'career\s*objective',
                r'profile',
                r'about',
                r'executive\s*summary'
            ],
            'experience': [
                r'work\s*experience',
                r'professional\s*experience',
                r'employment\s*history',
                r'career\s*history',
                r'work\s*history'
            ],
            'education': [
                r'education',
                r'academic\s*background',
                r'qualifications',
                r'degrees',
                r'academic\s*history'
            ],
            'skills': [
                r'technical\s*skills',
                r'skills',
                r'competencies',
                r'expertise',
                r'technologies',
                r'programming\s*languages'
            ],
            'projects': [
                r'projects',
                r'portfolio',
                r'achievements',
                r'accomplishments',
                r'key\s*projects'
            ],
            'certifications': [
                r'certifications',
                r'certificates',
                r'licenses',
                r'credentials',
                r'professional\s*certifications'
            ],
            'languages': [
                r'languages',
                r'language\s*skills',
                r'fluency',
                r'proficiency'
            ],
            'volunteer': [
                r'volunteer\s*work',
                r'volunteering',
                r'community\s*service',
                r'charity\s*work'
            ],
            'awards': [
                r'awards',
                r'honors',
                r'recognition',
                r'achievements'
            ]
        }
        
        # Enhanced patterns for information extraction
        self.extraction_patterns = {
            'name': [
                r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})$',
                r'^([A-Z][a-z]+\s+[A-Z]\.\s+[A-Z][a-z]+)$'
            ],
            'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            'phone': [
                r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
                r'\(\d{3}\)\s*\d{3}-\d{4}',
                r'\d{3}-\d{3}-\d{4}',
                r'\d{3}\.\d{3}\.\d{4}'
            ],
            'linkedin': r'linkedin\.com/in/[A-Za-z0-9-]+',
            'github': r'github\.com/[A-Za-z0-9-]+',
            'website': r'https?://[^\s]+',
            'date_range': [
                r'(\w+\s+\d{4})\s*[-–—]\s*(\w+\s+\d{4}|\bPresent\b)',
                r'(\d{4})\s*[-–—]\s*(\d{4}|\bPresent\b)',
                r'(\w+\s+\d{4})\s*to\s*(\w+\s+\d{4}|\bPresent\b)'
            ],
            'company': [
                r'([A-Z][A-Za-z\s&.,]+(?:Inc|Corp|LLC|Ltd|Company|Technologies|Solutions|Systems|Group|Enterprises|Industries))',
                r'([A-Z][A-Za-z\s&.,]+(?:University|College|Institute|School))'
            ],
            'job_title': [
                r'(Senior|Junior|Lead|Principal|Staff|Software|Data|Product|DevOps|Full\s*Stack|Frontend|Backend|Mobile|QA|Test|UI/UX|DevOps|Cloud|Security|Network|Database|System|Business|Project|Program|Product|Marketing|Sales|HR|Finance|Legal|Medical|Research|Teaching|Administrative)\s+(Engineer|Developer|Scientist|Manager|Analyst|Architect|Designer|Consultant|Coordinator|Specialist|Assistant|Director|Officer|Representative|Advisor|Consultant|Researcher|Instructor|Professor|Administrator)',
                r'(CEO|CTO|CFO|COO|VP|Director|Head|Team\s*Lead|Scrum\s*Master|Product\s*Owner)'
            ]
        }
        
        # Common skills and technologies
        self.skill_keywords = {
            'programming_languages': [
                'Python', 'JavaScript', 'Java', 'C++', 'C#', 'PHP', 'Ruby', 'Go', 'Rust', 'Swift', 'Kotlin', 'Scala', 'TypeScript'
            ],
            'frameworks': [
                'React', 'Angular', 'Vue', 'Node.js', 'Django', 'Flask', 'Spring', 'Express', 'Laravel', 'ASP.NET', 'Ruby on Rails'
            ],
            'databases': [
                'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'SQLite', 'Oracle', 'SQL Server', 'Cassandra', 'DynamoDB'
            ],
            'cloud_platforms': [
                'AWS', 'Azure', 'GCP', 'Heroku', 'DigitalOcean', 'Vercel', 'Netlify'
            ],
            'tools': [
                'Git', 'Docker', 'Kubernetes', 'Jenkins', 'GitLab', 'GitHub', 'Jira', 'Confluence', 'Slack', 'Trello'
            ],
            'methodologies': [
                'Agile', 'Scrum', 'Kanban', 'Waterfall', 'DevOps', 'CI/CD', 'TDD', 'BDD'
            ]
        }
    
    def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """
        Parse resume with improved accuracy and structure preservation.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Dictionary containing structured resume information
        """
        logger.info(f"Parsing resume: {file_path}")
        
        # Extract and clean text
        raw_text = self._extract_text(file_path)
        cleaned_text = self._clean_text(raw_text)
        
        # Parse using multiple strategies
        results = {}
        
        # Strategy 1: Section-based parsing
        section_results = self._parse_sections(cleaned_text)
        results.update(section_results)
        
        # Strategy 2: Pattern-based extraction
        pattern_results = self._extract_patterns(cleaned_text)
        results.update(pattern_results)
        
        # Strategy 3: Structured content parsing
        structured_results = self._parse_structured_content(cleaned_text)
        results.update(structured_results)
        
        # Strategy 4: Skills extraction
        skills_results = self._extract_skills(cleaned_text)
        results.update(skills_results)
        
        # Merge and validate results
        final_results = self._merge_results(results)
        final_results = self._validate_and_enhance(final_results, cleaned_text)
        
        # Add metadata
        final_results['metadata'] = {
            'parsed_at': datetime.now().isoformat(),
            'file_path': file_path,
            'file_size': os.path.getsize(file_path),
            'confidence_score': self._calculate_confidence(final_results)
        }
        
        logger.info("Resume parsing completed")
        return final_results
    
    def _extract_text(self, file_path: str) -> str:
        """Extract text from various file formats with better formatting preservation."""
        file_extension = file_path.lower().split('.')[-1]
        
        try:
            if file_extension == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension == 'docx':
                return self._extract_from_docx(file_path)
            elif file_extension == 'txt':
                return self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF with better formatting."""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    # Preserve line breaks and structure
                    text += f"\n--- PAGE {page_num + 1} ---\n{page_text}\n"
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            raise
        return text
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX with formatting preservation."""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text += row_text + "\n"
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            raise
        return text
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading TXT: {e}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text while preserving structure."""
        # Remove excessive whitespace but preserve line breaks
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:
                # Clean up common OCR issues
                line = re.sub(r'[|]', 'I', line)  # Fix common OCR mistake
                line = re.sub(r'[0]', 'O', line)  # Fix common OCR mistake
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def _parse_sections(self, text: str) -> Dict[str, Any]:
        """Parse resume by identifying and extracting sections."""
        lines = text.split('\n')
        sections = {}
        
        current_section = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if this line is a section header
            section_found = self._identify_section_header(line)
            
            if section_found:
                # Save previous section
                if current_section and current_content:
                    sections[current_section] = self._process_section_content(current_section, current_content)
                
                # Start new section
                current_section = section_found
                current_content = []
            elif current_section:
                current_content.append(line)
        
        # Save last section
        if current_section and current_content:
            sections[current_section] = self._process_section_content(current_section, current_content)
        
        return sections
    
    def _identify_section_header(self, line: str) -> Optional[str]:
        """Identify if a line is a section header."""
        line_lower = line.lower()
        
        for section, patterns in self.section_patterns.items():
            for pattern in patterns:
                if re.search(pattern, line_lower, re.IGNORECASE):
                    return section
        
        return None
    
    def _process_section_content(self, section: str, content: List[str]) -> Any:
        """Process section content based on section type."""
        if section == 'contact':
            return self._process_contact_section(content)
        elif section == 'experience':
            return self._process_experience_section(content)
        elif section == 'education':
            return self._process_education_section(content)
        elif section == 'skills':
            return self._process_skills_section(content)
        elif section == 'projects':
            return self._process_projects_section(content)
        else:
            return content
    
    def _process_contact_section(self, content: List[str]) -> Dict[str, str]:
        """Process contact section to extract structured information."""
        contact_info = {}
        text = ' '.join(content)
        
        # Extract email
        emails = re.findall(self.extraction_patterns['email'], text)
        if emails:
            contact_info['email'] = emails[0]
        
        # Extract phone
        for pattern in self.extraction_patterns['phone']:
            phones = re.findall(pattern, text)
            if phones:
                contact_info['phone'] = phones[0]
                break
        
        # Extract LinkedIn
        linkedin = re.findall(self.extraction_patterns['linkedin'], text)
        if linkedin:
            contact_info['linkedin'] = linkedin[0]
        
        # Extract GitHub
        github = re.findall(self.extraction_patterns['github'], text)
        if github:
            contact_info['github'] = github[0]
        
        # Extract website
        websites = re.findall(self.extraction_patterns['website'], text)
        if websites:
            contact_info['website'] = websites[0]
        
        return contact_info
    
    def _process_experience_section(self, content: List[str]) -> List[Dict[str, str]]:
        """Process experience section to extract job entries."""
        entries = []
        current_entry = {}
        
        for line in content:
            line = line.strip()
            
            # Look for job title
            for pattern in self.extraction_patterns['job_title']:
                matches = re.findall(pattern, line, re.IGNORECASE)
                if matches:
                    if current_entry:
                        entries.append(current_entry)
                    current_entry = {'title': line}
                    break
            
            # Look for company
            for pattern in self.extraction_patterns['company']:
                matches = re.findall(pattern, line, re.IGNORECASE)
                if matches:
                    if current_entry and 'title' in current_entry:
                        current_entry['company'] = line
                    break
            
            # Look for date range
            for pattern in self.extraction_patterns['date_range']:
                matches = re.findall(pattern, line, re.IGNORECASE)
                if matches:
                    if current_entry and 'title' in current_entry:
                        current_entry['dates'] = line
                    break
        
        if current_entry:
            entries.append(current_entry)
        
        return entries
    
    def _process_education_section(self, content: List[str]) -> List[Dict[str, str]]:
        """Process education section to extract education entries."""
        entries = []
        current_entry = {}
        
        for line in content:
            line = line.strip()
            
            # Look for degree patterns
            if re.search(r'\b(?:Bachelor|Master|PhD|Doctorate|Associate|Diploma|Certificate)\b', line, re.IGNORECASE):
                if current_entry:
                    entries.append(current_entry)
                current_entry = {'degree': line}
            
            # Look for institution patterns
            elif re.search(r'\b(?:University|College|Institute|School)\b', line, re.IGNORECASE):
                if current_entry and 'degree' in current_entry:
                    current_entry['institution'] = line
        
        if current_entry:
            entries.append(current_entry)
        
        return entries
    
    def _process_skills_section(self, content: List[str]) -> List[str]:
        """Process skills section to extract individual skills."""
        skills = []
        text = ' '.join(content)
        
        # Split by common delimiters
        skill_list = re.split(r'[,;•\n]', text)
        
        for skill in skill_list:
            skill = skill.strip()
            if skill and len(skill) > 2:
                skills.append(skill)
        
        return skills
    
    def _process_projects_section(self, content: List[str]) -> List[Dict[str, str]]:
        """Process projects section to extract project entries."""
        projects = []
        current_project = {}
        
        for line in content:
            line = line.strip()
            
            # Look for project name patterns
            if re.search(r'^[A-Z][A-Za-z\s]+(?:App|System|Platform|Tool|Dashboard|Website|API|Bot|Application)', line):
                if current_project:
                    projects.append(current_project)
                current_project = {'name': line}
            
            # Look for technologies used
            elif re.search(r'\b(?:Technologies|Tech|Tools|Stack|Built with|Using)\b', line, re.IGNORECASE):
                if current_project and 'name' in current_project:
                    current_project['technologies'] = line
        
        if current_project:
            projects.append(current_project)
        
        return projects
    
    def _extract_patterns(self, text: str) -> Dict[str, Any]:
        """Extract information using regex patterns."""
        results = {}
        
        # Extract name from first few lines
        lines = text.split('\n')
        for line in lines[:10]:
            for pattern in self.extraction_patterns['name']:
                matches = re.findall(pattern, line)
                if matches:
                    results['name'] = matches[0]
                    break
            if 'name' in results:
                break
        
        # Extract dates
        dates = []
        for pattern in self.extraction_patterns['date_range']:
            date_matches = re.findall(pattern, text, re.IGNORECASE)
            dates.extend(date_matches)
        results['dates'] = dates
        
        return results
    
    def _parse_structured_content(self, text: str) -> Dict[str, Any]:
        """Parse structured content like bullet points and lists."""
        results = {}
        
        lines = text.split('\n')
        
        # Extract bullet points
        bullet_points = []
        for line in lines:
            line = line.strip()
            if line.startswith(('•', '-', '*', '→', '▶', '○', '▪', '▫')):
                bullet_points.append(line[1:].strip())
        
        results['bullet_points'] = bullet_points
        
        return results
    
    def _extract_skills(self, text: str) -> Dict[str, List[str]]:
        """Extract skills using keyword matching."""
        skills = {}
        text_lower = text.lower()
        
        for category, skill_list in self.skill_keywords.items():
            found_skills = []
            for skill in skill_list:
                if skill.lower() in text_lower:
                    found_skills.append(skill)
            if found_skills:
                skills[category] = found_skills
        
        return skills
    
    def _merge_results(self, results: Dict[str, Any]) -> Dict[str, Any]:
        """Merge results from different parsing strategies."""
        merged = {
            'name': '',
            'email': '',
            'phone': '',
            'linkedin': '',
            'github': '',
            'website': '',
            'summary': '',
            'experience': [],
            'education': [],
            'skills': [],
            'projects': [],
            'certifications': [],
            'languages': [],
            'dates': [],
            'bullet_points': []
        }
        
        # Merge contact information
        if 'contact' in results:
            contact = results['contact']
            merged.update({k: v for k, v in contact.items() if v})
        
        # Merge name
        if 'name' in results:
            merged['name'] = results['name']
        
        # Merge sections
        for section in ['summary', 'experience', 'education', 'skills', 'projects', 'certifications', 'languages']:
            if section in results:
                merged[section] = results[section]
        
        # Merge dates
        if 'dates' in results:
            merged['dates'] = results['dates']
        
        # Merge bullet points
        if 'bullet_points' in results:
            merged['bullet_points'] = results['bullet_points']
        
        # Merge skills by category
        for key, value in results.items():
            if key in self.skill_keywords:
                merged[key] = value
        
        return merged
    
    def _validate_and_enhance(self, results: Dict[str, Any], text: str) -> Dict[str, Any]:
        """Validate and enhance the parsed results."""
        enhanced = results.copy()
        
        # Enhance name extraction if not found
        if not enhanced['name']:
            enhanced['name'] = self._enhance_name_extraction(text)
        
        # Enhance email extraction if not found
        if not enhanced['email']:
            enhanced['email'] = self._enhance_email_extraction(text)
        
        # Enhance skills if not found
        if not enhanced['skills']:
            enhanced['skills'] = self._enhance_skills_extraction(text)
        
        # Clean up empty entries
        for key, value in enhanced.items():
            if isinstance(value, list) and not value:
                enhanced[key] = []
            elif isinstance(value, str) and not value:
                enhanced[key] = ""
        
        return enhanced
    
    def _enhance_name_extraction(self, text: str) -> str:
        """Enhanced name extraction."""
        lines = text.split('\n')
        
        for line in lines[:5]:
            line = line.strip()
            if line and len(line.split()) <= 4:
                words = line.split()
                if all(word[0].isupper() for word in words if word):
                    return line
        
        return "Name not found"
    
    def _enhance_email_extraction(self, text: str) -> str:
        """Enhanced email extraction."""
        emails = re.findall(self.extraction_patterns['email'], text)
        return emails[0] if emails else "Email not found"
    
    def _enhance_skills_extraction(self, text: str) -> List[str]:
        """Enhanced skills extraction."""
        found_skills = []
        text_lower = text.lower()
        
        for category, skills in self.skill_keywords.items():
            for skill in skills:
                if skill.lower() in text_lower:
                    found_skills.append(skill)
        
        return found_skills if found_skills else ["Skills information not found"]
    
    def _calculate_confidence(self, results: Dict[str, Any]) -> int:
        """Calculate confidence score for the parsing results."""
        score = 0
        
        # Contact information (30 points)
        if results.get('name') and results['name'] != "Name not found":
            score += 10
        if results.get('email') and results['email'] != "Email not found":
            score += 10
        if results.get('phone'):
            score += 10
        
        # Experience (25 points)
        if results.get('experience'):
            score += 25
        
        # Education (20 points)
        if results.get('education'):
            score += 20
        
        # Skills (15 points)
        if results.get('skills') and results['skills'] != ["Skills information not found"]:
            score += 15
        
        # Additional sections (10 points)
        additional_sections = ['summary', 'projects', 'certifications', 'languages']
        for section in additional_sections:
            if results.get(section):
                score += 2.5
        
        return min(score, 100)
    
    def save_parsed_data(self, data: Dict[str, Any], output_file: str):
        """Save parsed data to JSON file."""
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            logger.info(f"Parsed data saved to {output_file}")
        except Exception as e:
            logger.error(f"Error saving parsed data: {e}")
    
    def get_parsing_report(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate a detailed parsing report."""
        report = {
            'confidence_score': data.get('metadata', {}).get('confidence_score', 0),
            'sections_found': [],
            'missing_sections': [],
            'data_quality': {},
            'recommendations': []
        }
        
        # Check which sections were found
        sections = ['summary', 'experience', 'education', 'skills', 'projects', 'certifications']
        for section in sections:
            if data.get(section):
                report['sections_found'].append(section)
            else:
                report['missing_sections'].append(section)
        
        # Data quality assessment
        if data.get('name') and data['name'] != "Name not found":
            report['data_quality']['name'] = "Found"
        else:
            report['data_quality']['name'] = "Missing"
            report['recommendations'].append("Name not found - check resume format")
        
        if data.get('email') and data['email'] != "Email not found":
            report['data_quality']['email'] = "Found"
        else:
            report['data_quality']['email'] = "Missing"
            report['recommendations'].append("Email not found - check resume format")
        
        if data.get('experience'):
            report['data_quality']['experience'] = f"Found {len(data['experience'])} entries"
        else:
            report['data_quality']['experience'] = "Missing"
            report['recommendations'].append("Work experience not found")
        
        if data.get('skills') and data['skills'] != ["Skills information not found"]:
            report['data_quality']['skills'] = f"Found {len(data['skills'])} skills"
        else:
            report['data_quality']['skills'] = "Missing"
            report['recommendations'].append("Skills section not found")
        
        return report 