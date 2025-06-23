import re
import PyPDF2
from docx import Document
import json
import spacy
from typing import Dict, List, Any, Optional
import logging
from datetime import datetime
import os

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class AdvancedResumeParser:
    def __init__(self):
        """Initialize the advanced resume parser with multiple parsing strategies."""
        self.sections = {
            'contact': ['contact', 'personal', 'info', 'details'],
            'summary': ['summary', 'objective', 'profile', 'about'],
            'experience': ['experience', 'work history', 'employment', 'work experience'],
            'education': ['education', 'academic', 'qualifications', 'degree'],
            'skills': ['skills', 'technical skills', 'competencies', 'expertise'],
            'projects': ['projects', 'portfolio', 'achievements', 'accomplishments'],
            'certifications': ['certifications', 'certificates', 'licenses', 'credentials'],
            'languages': ['languages', 'language skills', 'fluency'],
            'volunteer': ['volunteer', 'volunteering', 'community service'],
            'awards': ['awards', 'honors', 'recognition', 'achievements']
        }
        
        # Enhanced patterns for better extraction
        self.patterns = {
            'name': [
                r'^[A-Z][a-z]+ [A-Z][a-z]+$',
                r'^[A-Z][a-z]+ [A-Z]\. [A-Z][a-z]+$',
                r'^[A-Z][a-z]+ [A-Z][a-z]+ [A-Z][a-z]+$'
            ],
            'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            'phone': [
                r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
                r'\(\d{3}\) \d{3}-\d{4}',
                r'\d{3}-\d{3}-\d{4}'
            ],
            'linkedin': r'linkedin\.com/in/[A-Za-z0-9-]+',
            'github': r'github\.com/[A-Za-z0-9-]+',
            'website': r'https?://[^\s]+',
            'date': r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{4}\b',
            'year': r'\b(19|20)\d{2}\b'
        }
        
        # Try to load spaCy model for NLP
        try:
            self.nlp = spacy.load("en_core_web_sm")
            self.use_nlp = True
        except OSError:
            logger.warning("spaCy model not found. Install with: python -m spacy download en_core_web_sm")
            self.use_nlp = False
    
    def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """
        Parse resume using multiple strategies for better accuracy.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Dictionary containing structured resume information
        """
        logger.info(f"Parsing resume: {file_path}")
        
        # Extract text from file
        text = self._extract_text(file_path)
        if not text:
            raise ValueError(f"Could not extract text from {file_path}")
        
        # Parse using multiple strategies
        results = {}
        
        # Strategy 1: Section-based parsing
        section_results = self._parse_by_sections(text)
        results.update(section_results)
        
        # Strategy 2: Pattern-based parsing
        pattern_results = self._parse_by_patterns(text)
        results.update(pattern_results)
        
        # Strategy 3: NLP-based parsing (if available)
        if self.use_nlp:
            nlp_results = self._parse_with_nlp(text)
            results.update(nlp_results)
        
        # Strategy 4: Structured parsing
        structured_results = self._parse_structured_content(text)
        results.update(structured_results)
        
        # Merge and clean results
        final_results = self._merge_and_clean_results(results)
        
        # Validate and enhance results
        final_results = self._validate_and_enhance(final_results, text)
        
        logger.info("Resume parsing completed")
        return final_results
    
    def _extract_text(self, file_path: str) -> str:
        """Extract text from various file formats."""
        file_extension = file_path.lower().split('.')[-1]
        
        try:
            if file_extension == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension == 'docx':
                return self._extract_from_docx(file_path)
            elif file_extension == 'txt':
                return self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF with better formatting preservation."""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    # Clean up common PDF extraction issues
                    page_text = self._clean_pdf_text(page_text)
                    text += f"\n--- PAGE {page_num + 1} ---\n{page_text}\n"
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            raise
        return text
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX with formatting preservation."""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract from tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells if cell.text.strip()])
                    if row_text.strip():
                        text += row_text + "\n"
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            raise
        return text
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading TXT: {e}")
            raise
    
    def _clean_pdf_text(self, text: str) -> str:
        """Clean up common PDF extraction issues."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Fix common OCR issues
        text = re.sub(r'[|]', 'I', text)  # Fix common OCR mistake
        text = re.sub(r'[0]', 'O', text)  # Fix common OCR mistake
        
        # Remove page numbers and headers
        text = re.sub(r'\b\d+\s*of\s*\d+\b', '', text)
        
        return text.strip()
    
    def _parse_by_sections(self, text: str) -> Dict[str, Any]:
        """Parse resume by identifying and extracting sections."""
        lines = text.split('\n')
        sections = {}
        
        current_section = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if this line is a section header
            section_found = self._identify_section_header(line)
            
            if section_found:
                # Save previous section
                if current_section and current_content:
                    sections[current_section] = self._clean_section_content(current_content)
                
                # Start new section
                current_section = section_found
                current_content = []
            elif current_section:
                current_content.append(line)
        
        # Save last section
        if current_section and current_content:
            sections[current_section] = self._clean_section_content(current_content)
        
        return sections
    
    def _identify_section_header(self, line: str) -> Optional[str]:
        """Identify if a line is a section header."""
        line_lower = line.lower()
        
        for section, keywords in self.sections.items():
            for keyword in keywords:
                if keyword in line_lower and len(line.split()) <= 3:
                    return section
        
        return None
    
    def _clean_section_content(self, content: List[str]) -> List[str]:
        """Clean and structure section content."""
        cleaned = []
        for item in content:
            item = item.strip()
            if item and len(item) > 2:
                cleaned.append(item)
        return cleaned
    
    def _parse_by_patterns(self, text: str) -> Dict[str, Any]:
        """Extract information using regex patterns."""
        results = {}
        
        # Extract contact information
        results['contact'] = self._extract_contact_info(text)
        
        # Extract dates and years
        results['dates'] = self._extract_dates(text)
        
        # Extract URLs
        results['urls'] = self._extract_urls(text)
        
        return results
    
    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information using patterns."""
        contact = {}
        
        # Extract name
        for pattern in self.patterns['name']:
            matches = re.findall(pattern, text, re.MULTILINE)
            if matches:
                contact['name'] = matches[0]
                break
        
        # Extract email
        emails = re.findall(self.patterns['email'], text)
        if emails:
            contact['email'] = emails[0]
        
        # Extract phone
        for pattern in self.patterns['phone']:
            phones = re.findall(pattern, text)
            if phones:
                contact['phone'] = phones[0]
                break
        
        # Extract LinkedIn
        linkedin = re.findall(self.patterns['linkedin'], text)
        if linkedin:
            contact['linkedin'] = linkedin[0]
        
        # Extract GitHub
        github = re.findall(self.patterns['github'], text)
        if github:
            contact['github'] = github[0]
        
        return contact
    
    def _extract_dates(self, text: str) -> List[str]:
        """Extract dates from text."""
        dates = []
        
        # Extract full dates
        full_dates = re.findall(self.patterns['date'], text)
        dates.extend(full_dates)
        
        # Extract years
        years = re.findall(self.patterns['year'], text)
        dates.extend(years)
        
        return list(set(dates))
    
    def _extract_urls(self, text: str) -> List[str]:
        """Extract URLs from text."""
        urls = re.findall(self.patterns['website'], text)
        return list(set(urls))
    
    def _parse_with_nlp(self, text: str) -> Dict[str, Any]:
        """Parse resume using NLP techniques."""
        if not self.use_nlp:
            return {}
        
        doc = self.nlp(text)
        results = {}
        
        # Extract named entities
        entities = {}
        for ent in doc.ents:
            if ent.label_ not in entities:
                entities[ent.label_] = []
            entities[ent.label_].append(ent.text)
        
        results['entities'] = entities
        
        # Extract organizations (companies)
        if 'ORG' in entities:
            results['companies'] = list(set(entities['ORG']))
        
        # Extract locations
        if 'GPE' in entities:
            results['locations'] = list(set(entities['GPE']))
        
        # Extract dates
        if 'DATE' in entities:
            results['nlp_dates'] = list(set(entities['DATE']))
        
        return results
    
    def _parse_structured_content(self, text: str) -> Dict[str, Any]:
        """Parse structured content like bullet points and lists."""
        results = {}
        
        lines = text.split('\n')
        
        # Extract bullet points
        bullet_points = []
        for line in lines:
            line = line.strip()
            if line.startswith(('•', '-', '*', '→', '▶')):
                bullet_points.append(line[1:].strip())
        
        results['bullet_points'] = bullet_points
        
        # Extract experience entries (look for date patterns)
        experience_entries = self._extract_experience_entries(lines)
        results['experience_entries'] = experience_entries
        
        # Extract education entries
        education_entries = self._extract_education_entries(lines)
        results['education_entries'] = education_entries
        
        return results
    
    def _extract_experience_entries(self, lines: List[str]) -> List[Dict[str, str]]:
        """Extract structured experience entries."""
        entries = []
        current_entry = {}
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            # Look for job title patterns
            if re.search(r'\b(?:Senior|Junior|Lead|Principal|Staff|Software|Data|Product|DevOps|Full Stack|Frontend|Backend)\s+(?:Engineer|Developer|Scientist|Manager|Analyst|Architect)\b', line, re.IGNORECASE):
                if current_entry:
                    entries.append(current_entry)
                current_entry = {'title': line}
            
            # Look for company patterns
            elif re.search(r'\b(?:Inc|Corp|LLC|Ltd|Company|Technologies|Solutions|Systems|Group)\b', line, re.IGNORECASE):
                if current_entry and 'title' in current_entry:
                    current_entry['company'] = line
            
            # Look for date patterns
            elif re.search(r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{4}\b', line):
                if current_entry and 'title' in current_entry:
                    current_entry['dates'] = line
        
        if current_entry:
            entries.append(current_entry)
        
        return entries
    
    def _extract_education_entries(self, lines: List[str]) -> List[Dict[str, str]]:
        """Extract structured education entries."""
        entries = []
        current_entry = {}
        
        for line in lines:
            line = line.strip()
            
            # Look for degree patterns
            if re.search(r'\b(?:Bachelor|Master|PhD|Doctorate|Associate|Diploma|Certificate)\b', line, re.IGNORECASE):
                if current_entry:
                    entries.append(current_entry)
                current_entry = {'degree': line}
            
            # Look for university patterns
            elif re.search(r'\b(?:University|College|Institute|School)\b', line, re.IGNORECASE):
                if current_entry and 'degree' in current_entry:
                    current_entry['institution'] = line
        
        if current_entry:
            entries.append(current_entry)
        
        return entries
    
    def _merge_and_clean_results(self, results: Dict[str, Any]) -> Dict[str, Any]:
        """Merge results from different parsing strategies."""
        merged = {
            'name': '',
            'email': '',
            'phone': '',
            'linkedin': '',
            'github': '',
            'summary': '',
            'experience': [],
            'education': [],
            'skills': [],
            'projects': [],
            'certifications': [],
            'languages': [],
            'companies': [],
            'locations': [],
            'dates': []
        }
        
        # Merge contact information
        if 'contact' in results:
            contact = results['contact']
            merged.update({k: v for k, v in contact.items() if v})
        
        # Merge sections
        for section in ['summary', 'experience', 'education', 'skills', 'projects', 'certifications', 'languages']:
            if section in results:
                merged[section] = results[section]
        
        # Merge entities
        if 'companies' in results:
            merged['companies'] = results['companies']
        if 'locations' in results:
            merged['locations'] = results['locations']
        if 'dates' in results:
            merged['dates'] = results['dates']
        
        # Merge structured content
        if 'experience_entries' in results:
            merged['experience'].extend(results['experience_entries'])
        if 'education_entries' in results:
            merged['education'].extend(results['education_entries'])
        
        return merged
    
    def _validate_and_enhance(self, results: Dict[str, Any], original_text: str) -> Dict[str, Any]:
        """Validate and enhance the parsed results."""
        enhanced = results.copy()
        
        # Enhance name extraction if not found
        if not enhanced['name'] or enhanced['name'] == 'Name not found':
            enhanced['name'] = self._enhance_name_extraction(original_text)
        
        # Enhance email extraction if not found
        if not enhanced['email'] or enhanced['email'] == 'Email not found':
            enhanced['email'] = self._enhance_email_extraction(original_text)
        
        # Enhance skills extraction
        if not enhanced['skills'] or enhanced['skills'] == ['Skills information not found']:
            enhanced['skills'] = self._enhance_skills_extraction(original_text)
        
        # Clean up empty or invalid entries
        for key, value in enhanced.items():
            if isinstance(value, list) and value == [f"{key.title()} information not found"]:
                enhanced[key] = []
            elif isinstance(value, str) and value == f"{key.title()} not found":
                enhanced[key] = ""
        
        return enhanced
    
    def _enhance_name_extraction(self, text: str) -> str:
        """Enhanced name extraction using multiple strategies."""
        lines = text.split('\n')
        
        # Look for name in first few lines
        for line in lines[:5]:
            line = line.strip()
            if line and len(line.split()) <= 4:
                # Check if it looks like a name
                words = line.split()
                if all(word[0].isupper() for word in words if word):
                    return line
        
        return "Name not found"
    
    def _enhance_email_extraction(self, text: str) -> str:
        """Enhanced email extraction."""
        emails = re.findall(self.patterns['email'], text)
        return emails[0] if emails else "Email not found"
    
    def _enhance_skills_extraction(self, text: str) -> List[str]:
        """Enhanced skills extraction using keyword matching."""
        # Common technical skills
        skill_keywords = [
            'Python', 'JavaScript', 'Java', 'C++', 'C#', 'PHP', 'Ruby', 'Go', 'Rust',
            'HTML', 'CSS', 'SQL', 'NoSQL', 'MongoDB', 'PostgreSQL', 'MySQL',
            'React', 'Angular', 'Vue', 'Node.js', 'Django', 'Flask', 'Spring',
            'AWS', 'Azure', 'GCP', 'Docker', 'Kubernetes', 'Git', 'Jenkins',
            'Machine Learning', 'AI', 'Data Science', 'Statistics', 'R',
            'Tableau', 'Power BI', 'Excel', 'PowerPoint', 'Word'
        ]
        
        found_skills = []
        text_lower = text.lower()
        
        for skill in skill_keywords:
            if skill.lower() in text_lower:
                found_skills.append(skill)
        
        return found_skills if found_skills else ["Skills information not found"]
    
    def save_parsed_data(self, data: Dict[str, Any], output_file: str):
        """Save parsed data to JSON file."""
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            logger.info(f"Parsed data saved to {output_file}")
        except Exception as e:
            logger.error(f"Error saving parsed data: {e}")
    
    def get_parsing_stats(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Get statistics about the parsing results."""
        stats = {
            'total_sections_found': 0,
            'contact_info_complete': False,
            'has_experience': False,
            'has_education': False,
            'has_skills': False,
            'confidence_score': 0
        }
        
        # Count sections
        for section in ['summary', 'experience', 'education', 'skills', 'projects', 'certifications']:
            if data.get(section) and data[section] != [] and data[section] != "":
                stats['total_sections_found'] += 1
        
        # Check contact info
        contact_fields = ['name', 'email', 'phone']
        stats['contact_info_complete'] = all(data.get(field) for field in contact_fields)
        
        # Check other sections
        stats['has_experience'] = bool(data.get('experience'))
        stats['has_education'] = bool(data.get('education'))
        stats['has_skills'] = bool(data.get('skills'))
        
        # Calculate confidence score
        score = 0
        if stats['contact_info_complete']:
            score += 30
        if stats['has_experience']:
            score += 25
        if stats['has_education']:
            score += 20
        if stats['has_skills']:
            score += 15
        score += min(stats['total_sections_found'] * 5, 10)
        
        stats['confidence_score'] = score
        
        return stats 